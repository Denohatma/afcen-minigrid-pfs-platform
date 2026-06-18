"""Professional DOCX writer for AfCEN Minigrid PFS documents.

Converts rendered markdown into a styled Word document with:
- Professional cover page with branded horizontal rules
- Smart key-value pair rendering (borderless info blocks, not grid tables)
- Properly styled data tables with branded navy header rows
- Embedded chart images and location maps
- Page breaks between major sections
- Headers with site name, footers with page numbers
"""

from __future__ import annotations

import logging
import re
from datetime import date
from pathlib import Path

logger = logging.getLogger("minigrid")

_NAVY = "1B3A5C"
_ACCENT = "2E75B6"
_DARK = "2D2D2D"
_GRAY = "666666"
_LIGHT_GRAY = "999999"
_ROW_ALT = "F5F6F8"
_BORDER = "D5D5D5"
_WHITE = "FFFFFF"
_RED = "C0392B"
_GREEN = "27AE60"
_AMBER = "F39C12"
_FONT = "Calibri"

_KV_FIRST = {"parameter", "metric", "component", "factor", "milestone",
             "summary metric", "input", "source", "risk category",
             "loss category"}
_KV_SECOND = {"value", "specification", "assessment", "target", "provenance",
              "loss"}

_STATUS_COLORS = {
    "\u2713": _GREEN, "\u2713 OK": _GREEN,
    "\u2717": _RED, "\u2717 Failed": _RED,
    "\u25B3": _AMBER,
}

# Lazy-loaded docx imports — set by _load_docx()
_Pt = _Cm = _RGBColor = _Inches = None
_WD_ALIGN = _WD_TABLE_ALIGN = _WD_LINE_SPACING = None
_OxmlElement = _parse_xml = _qn = _nsdecls = None

_RGB_CACHE: dict[str, object] = {}


def _load_docx():
    """Import all python-docx symbols once; raises ImportError if missing."""
    global _Pt, _Cm, _RGBColor, _Inches
    global _WD_ALIGN, _WD_TABLE_ALIGN, _WD_LINE_SPACING
    global _OxmlElement, _parse_xml, _qn, _nsdecls

    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import OxmlElement, parse_xml

    _Pt, _Cm, _RGBColor, _Inches = Pt, Cm, RGBColor, Inches
    _WD_ALIGN = WD_ALIGN_PARAGRAPH
    _WD_TABLE_ALIGN = WD_TABLE_ALIGNMENT
    _WD_LINE_SPACING = WD_LINE_SPACING
    _OxmlElement, _parse_xml, _qn, _nsdecls = OxmlElement, parse_xml, qn, nsdecls


def _rgb(h):
    cached = _RGB_CACHE.get(h)
    if cached is None:
        cached = _RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
        _RGB_CACHE[h] = cached
    return cached


def _strip_md(text):
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"`(.*?)`", r"\1", text)
    return text


# ── Entry point ────────────────────────────────────────────────────────

def write_docx(markdown_text: str, output_path: str, site_data=None,
               chart_paths: dict | None = None):
    try:
        _load_docx()
        from docx import Document
    except ImportError:
        logger.warning("python-docx not installed. Skipping DOCX output.")
        return

    doc = Document()
    _setup_styles(doc)
    _setup_page(doc)
    _add_cover_page(doc, site_data)
    doc.add_page_break()
    _add_toc(doc)
    doc.add_page_break()
    _render_body(doc, markdown_text.split("\n"), chart_paths or {})
    _add_header_footer(doc, site_data)
    doc.save(output_path)
    logger.info(f"DOCX saved to {output_path}")


# ── Styles ─────────────────────────────────────────────────────────────

def _setup_styles(doc):
    s = doc.styles["Normal"]
    s.font.name = _FONT
    s.font.size = _Pt(10)
    s.font.color.rgb = _rgb(_DARK)
    pf = s.paragraph_format
    pf.space_before = _Pt(2)
    pf.space_after = _Pt(6)
    pf.line_spacing_rule = _WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15

    for level, size, color, before, after in [
        (1, 18, _NAVY, 0, 10),
        (2, 13, _ACCENT, 16, 6),
        (3, 11, _NAVY, 12, 4),
    ]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = _FONT
        h.font.size = _Pt(size)
        h.font.bold = True
        h.font.color.rgb = _rgb(color)
        h.paragraph_format.space_before = _Pt(before)
        h.paragraph_format.space_after = _Pt(after)
        h.paragraph_format.keep_with_next = True


def _setup_page(doc):
    for sec in doc.sections:
        sec.top_margin = _Cm(2.0)
        sec.bottom_margin = _Cm(2.0)
        sec.left_margin = _Cm(2.54)
        sec.right_margin = _Cm(2.54)


# ── Cover page ─────────────────────────────────────────────────────────

def _add_cover_page(doc, site_data):
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = _WD_ALIGN.CENTER
    run = p.add_run("PRE-FEASIBILITY STUDY")
    run.font.size = _Pt(32)
    run.font.color.rgb = _rgb(_NAVY)
    run.bold = True
    run.font.name = _FONT

    doc.add_paragraph()
    _add_rule(doc)
    doc.add_paragraph()

    if not site_data:
        return

    p = doc.add_paragraph()
    p.alignment = _WD_ALIGN.CENTER
    run = p.add_run(f"{site_data.site_name} Solar Hybrid Mini-Grid")
    run.font.size = _Pt(20)
    run.font.color.rgb = _rgb(_NAVY)
    run.font.name = _FONT

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = _WD_ALIGN.CENTER
    run = p.add_run(
        f"{site_data.district} LGA, {site_data.province} State, "
        f"{site_data.country}"
    )
    run.font.size = _Pt(13)
    run.font.color.rgb = _rgb(_GRAY)
    run.font.name = _FONT

    for _ in range(3):
        doc.add_paragraph()

    if site_data.developer:
        p = doc.add_paragraph()
        p.alignment = _WD_ALIGN.CENTER
        run = p.add_run(f"Prepared for: {site_data.developer}")
        run.font.size = _Pt(11)
        run.font.name = _FONT

    p = doc.add_paragraph()
    p.alignment = _WD_ALIGN.CENTER
    run = p.add_run(date.today().strftime("%B %Y"))
    run.font.size = _Pt(11)
    run.font.name = _FONT

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = _WD_ALIGN.CENTER
    run = p.add_run("Generated by AfCEN Minigrid PFS Platform v3.0")
    run.font.size = _Pt(9)
    run.font.color.rgb = _rgb(_GRAY)
    run.font.name = _FONT

    doc.add_paragraph()
    _add_rule(doc)

    p = doc.add_paragraph()
    p.alignment = _WD_ALIGN.CENTER
    run = p.add_run(
        "CONFIDENTIAL \u2014 This document contains proprietary information "
        "prepared for the exclusive use of the intended recipient."
    )
    run.font.size = _Pt(8)
    run.font.color.rgb = _rgb(_LIGHT_GRAY)
    run.italic = True
    run.font.name = _FONT


def _add_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = _OxmlElement("w:pBdr")
    bottom = _OxmlElement("w:bottom")
    bottom.set(_qn("w:val"), "single")
    bottom.set(_qn("w:sz"), "12")
    bottom.set(_qn("w:color"), _ACCENT)
    bottom.set(_qn("w:space"), "1")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Table of contents ──────────────────────────────────────────────────

def _add_toc(doc):
    doc.add_heading("Table of Contents", level=1)
    p = doc.add_paragraph()
    run = p.add_run(
        "Update this table of contents after opening in Microsoft Word: "
        "select all (Ctrl+A), then press F9."
    )
    run.font.size = _Pt(9)
    run.font.color.rgb = _rgb(_GRAY)
    run.italic = True
    run.font.name = _FONT


# ── Image embedding ───────────────────────────────────────────────────

# Map chart keys → section heading text fragments where they should appear
_CHART_PLACEMENTS = {
    "location_map": "3. Site Overview",
    "load_profile": "4. Demand Assessment",
    "solar_resource": "5. Solar Resource",
    "monthly_gen_demand": "6. Load\u2013Generation Balance",
    "dispatch": "6. Load\u2013Generation Balance",
    "capex": "10. Financial Model",
    "cashflow": "10. Financial Model",
}

_CHART_CAPTIONS = {
    "location_map": "Figure: Site location",
    "load_profile": "Figure: Average 24-hour load profile",
    "solar_resource": "Figure: Monthly solar irradiance and temperature",
    "monthly_gen_demand": "Figure: Monthly generation vs. demand",
    "dispatch": "Figure: Sample day dispatch profile",
    "capex": "Figure: CAPEX breakdown by component",
    "cashflow": "Figure: 25-year cumulative cash flow",
}


def _insert_chart(doc, chart_path, caption):
    if not Path(chart_path).exists():
        return
    p = doc.add_paragraph()
    p.alignment = _WD_ALIGN.CENTER
    run = p.add_run()
    run.add_picture(chart_path, width=_Inches(5.8))

    cap = doc.add_paragraph()
    cap.alignment = _WD_ALIGN.CENTER
    run = cap.add_run(caption)
    run.font.size = _Pt(8)
    run.font.color.rgb = _rgb(_GRAY)
    run.italic = True
    run.font.name = _FONT


def _pending_charts_for_section(section_title, chart_paths):
    """Return list of (key, path, caption) for charts that belong after this heading."""
    results = []
    for key, fragment in _CHART_PLACEMENTS.items():
        if key in chart_paths and fragment.lower() in section_title.lower():
            results.append((key, chart_paths[key], _CHART_CAPTIONS.get(key, "")))
    return results


# ── Body rendering ─────────────────────────────────────────────────────

def _render_body(doc, lines, chart_paths):
    i = 0
    while i < len(lines):
        if lines[i].strip().startswith("## "):
            break
        i += 1

    first_section = True
    current_section = ""
    pending_charts = []

    while i < len(lines):
        stripped = lines[i].strip()

        if re.match(r"^-{3,}$", stripped):
            i += 1
            continue

        if stripped.startswith("## "):
            if not first_section:
                doc.add_page_break()
            first_section = False
            current_section = _strip_md(stripped[3:])
            doc.add_heading(current_section, level=1)
            pending_charts = _pending_charts_for_section(current_section, chart_paths)
            i += 1
            continue

        if stripped.startswith("#### "):
            doc.add_heading(_strip_md(stripped[5:]), level=3)
            i += 1
            continue

        if stripped.startswith("### "):
            # Insert pending charts before second subsection of their parent
            if pending_charts:
                for key, path, caption in pending_charts:
                    _insert_chart(doc, path, caption)
                pending_charts = []
            doc.add_heading(_strip_md(stripped[4:]), level=2)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            _process_table(doc, table_lines)
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            para = doc.add_paragraph(style="List Bullet")
            _add_rich_text(para, stripped[2:])
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            para = doc.add_paragraph(style="List Number")
            _add_rich_text(para, text)
            i += 1
            continue

        if stripped.startswith("[DATA GAP:"):
            para = doc.add_paragraph()
            para.paragraph_format.space_before = _Pt(6)
            para.paragraph_format.space_after = _Pt(6)
            run = para.add_run(stripped)
            run.font.color.rgb = _rgb(_RED)
            run.bold = True
            run.font.size = _Pt(10)
            run.font.name = _FONT
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        para = doc.add_paragraph()
        _add_rich_text(para, stripped)
        i += 1

    # Flush any remaining charts
    for key, path, caption in pending_charts:
        _insert_chart(doc, path, caption)


# ── Table processing ──────────────────────────────────────────────────

def _process_table(doc, table_lines):
    if len(table_lines) < 2:
        return

    headers = [c.strip() for c in table_lines[0].split("|") if c.strip()]
    n_cols = len(headers)
    if n_cols == 0:
        return

    data = []
    for line in table_lines[2:]:
        cells = [c.strip() for c in line.split("|") if c.strip()]
        if cells:
            while len(cells) < n_cols:
                cells.append("")
            data.append(cells[:n_cols])

    if not data:
        return

    if _is_kv_table(headers, data):
        _add_kv_block(doc, data)
    else:
        _add_data_table(doc, headers, data)


def _is_kv_table(headers, data):
    if len(headers) != 2 or len(data) > 15:
        return False
    h1, h2 = headers[0].lower().strip(), headers[1].lower().strip()
    return (any(kw in h1 for kw in _KV_FIRST) or
            any(kw in h2 for kw in _KV_SECOND))


def _add_kv_block(doc, data):
    table = doc.add_table(rows=len(data), cols=2)
    table.alignment = _WD_TABLE_ALIGN.LEFT

    for i, (label, value) in enumerate(data):
        row = table.rows[i]
        lc = row.cells[0]
        lc.text = ""
        lc.width = _Cm(6.5)
        p = lc.paragraphs[0]
        p.paragraph_format.space_before = _Pt(3)
        p.paragraph_format.space_after = _Pt(3)
        run = p.add_run(_strip_md(label))
        run.font.size = _Pt(9.5)
        run.font.name = _FONT
        run.font.color.rgb = _rgb(_GRAY)
        run.bold = True

        vc = row.cells[1]
        vc.text = ""
        vc.width = _Cm(9.5)
        p = vc.paragraphs[0]
        p.paragraph_format.space_before = _Pt(3)
        p.paragraph_format.space_after = _Pt(3)
        _add_rich_text(p, value, size_pt=10)

        if i % 2 == 0:
            _shade_cell(lc, _ROW_ALT)
            _shade_cell(vc, _ROW_ALT)

    _set_borders(
        table,
        top=("single", 6, _ACCENT),
        bottom=("single", 6, _ACCENT),
        insideH=("single", 2, _BORDER),
    )


def _add_data_table(doc, headers, data):
    n_cols = len(headers)
    fsz = 9 if n_cols <= 5 else 8

    table = doc.add_table(rows=1 + len(data), cols=n_cols)
    table.alignment = _WD_TABLE_ALIGN.LEFT

    for j, ht in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        _shade_cell(cell, _NAVY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = _Pt(4)
        p.paragraph_format.space_after = _Pt(4)
        run = p.add_run(_strip_md(ht))
        run.bold = True
        run.font.size = _Pt(fsz)
        run.font.name = _FONT
        run.font.color.rgb = _rgb(_WHITE)

    for i, row_data in enumerate(data):
        is_summary = row_data[0].strip().startswith("**")
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = _Pt(2)
            p.paragraph_format.space_after = _Pt(2)
            _add_rich_text(p, cell_text, size_pt=fsz)

            clean = _strip_md(cell_text).strip()
            if clean in _STATUS_COLORS:
                c = _STATUS_COLORS[clean]
                for run in p.runs:
                    run.font.color.rgb = _rgb(c)
                    run.bold = True

            if not is_summary and i % 2 == 1:
                _shade_cell(cell, _ROW_ALT)

    _set_borders(
        table,
        top=("single", 8, _NAVY),
        bottom=("single", 8, _NAVY),
        insideH=("single", 2, _BORDER),
    )


# ── Rich text ──────────────────────────────────────────────────────────

def _add_rich_text(para, text, size_pt=10):
    parts = re.split(r"(\*\*.*?\*\*|\*[^*]+?\*)", text)
    pt = _Pt(size_pt)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            run = para.add_run(part)
        run.font.size = pt
        run.font.name = _FONT


# ── XML helpers ────────────────────────────────────────────────────────

def _shade_cell(cell, hex_color):
    shading = _parse_xml(
        f'<w:shd {_nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _set_borders(table, *, top=None, bottom=None, left=None, right=None,
                 insideH=None, insideV=None):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = _OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for existing in tblPr.findall(_qn("w:tblBorders")):
        tblPr.remove(existing)

    borders = _OxmlElement("w:tblBorders")
    for edge, spec in {"top": top, "left": left, "bottom": bottom,
                       "right": right, "insideH": insideH,
                       "insideV": insideV}.items():
        el = _OxmlElement(f"w:{edge}")
        if spec:
            el.set(_qn("w:val"), spec[0])
            el.set(_qn("w:sz"), str(spec[1]))
            el.set(_qn("w:color"), spec[2])
        else:
            el.set(_qn("w:val"), "none")
            el.set(_qn("w:sz"), "0")
        el.set(_qn("w:space"), "0")
        borders.append(el)
    tblPr.append(borders)


# ── Header and footer ─────────────────────────────────────────────────

def _add_header_footer(doc, site_data):
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = _WD_ALIGN.RIGHT
    label = (f"Pre-Feasibility Study \u2014 {site_data.site_name}"
             if site_data else "Pre-Feasibility Study")
    run = hp.add_run(label)
    run.font.size = _Pt(8)
    run.font.color.rgb = _rgb(_GRAY)
    run.font.name = _FONT
    run.italic = True

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = _WD_ALIGN.CENTER

    run = fp.add_run("Page ")
    run.font.size = _Pt(8)
    run.font.color.rgb = _rgb(_GRAY)
    run.font.name = _FONT

    r1 = fp.add_run()
    fld_begin = _OxmlElement("w:fldChar")
    fld_begin.set(_qn("w:fldCharType"), "begin")
    r1._r.append(fld_begin)

    r2 = fp.add_run()
    instr = _OxmlElement("w:instrText")
    instr.set(_qn("xml:space"), "preserve")
    instr.text = " PAGE "
    r2._r.append(instr)

    r3 = fp.add_run()
    fld_end = _OxmlElement("w:fldChar")
    fld_end.set(_qn("w:fldCharType"), "end")
    r3._r.append(fld_end)
